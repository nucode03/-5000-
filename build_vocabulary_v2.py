from __future__ import annotations

import csv
import re
from collections import defaultdict
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
CSV_PATH = ROOT / "build" / "oxford_5000_entries.csv"
OUT = ROOT / "output"
DOCX_PATH = OUT / "Oxford_5000_B2-C1_Korean_Vocabulary_ver2.docx"
PDF_PATH = OUT / "Oxford_5000_B2-C1_Korean_Vocabulary_ver2.pdf"
FONT = Path(r"C:\Windows\Fonts\malgun.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\malgunbd.ttf")


def korean_pos(source_detail: str) -> str:
    labels = {
        "n.": "명사", "v.": "동사", "adj.": "형용사", "adv.": "부사",
        "prep.": "전치사", "conj.": "접속사", "pron.": "대명사", "det.": "한정사",
        "exclam.": "감탄사", "number": "수사",
    }
    raw = re.sub(r"\b(?:B1|B2|C1)\b", "", source_detail)
    tokens = re.findall(r"n\.|v\.|adj\.|adv\.|prep\.|conj\.|pron\.|det\.|exclam\.|number", raw)
    if not tokens:
        raise ValueError(f"Unrecognized part of speech: {source_detail}")
    return " / ".join(labels[token] for token in tokens)


def load_entries():
    with CSV_PATH.open(encoding="utf-8-sig", newline="") as f:
        entries = list(csv.DictReader(f))
    for entry in entries:
        entry["pos_en"] = korean_pos(entry["source_detail"])
    if len(entries) != 2000 or any(not x["meaning"].strip() for x in entries):
        raise ValueError("The v2 source data must contain all 2,000 translated entries.")
    return entries


def set_docx_font(run, size: float, bold=False, color: str | None = None) -> None:
    run.font.name = "Malgun Gothic"
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_docx_text(p, text, size=9.4, bold=False, color=None, align=None) -> None:
    p.clear()
    set_docx_font(p.add_run(text), size, bold, color)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def repeat_header(row):
    prop = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    prop.append(node)


def page_field(p):
    run = p.add_run("Page ")
    set_docx_font(run, 8.5, color="687385")
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def make_docx(entries):
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = section.bottom_margin = Mm(14)
    section.left_margin = section.right_margin = Mm(14)
    section.header_distance = section.footer_distance = Mm(7)
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    for name, size, color in (("Heading 1", 19, "163A5F"), ("Heading 2", 13, "2E74B5")):
        style = doc.styles[name]
        style.font.name = "Malgun Gothic"; style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
    set_docx_text(section.header.paragraphs[0], "Oxford 5000 B2-C1 | Korean Vocabulary Book ver2", 8.5, color="687385")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_field(footer)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(124); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_docx_text(p, "Oxford 5000", 30, True, "163A5F", WD_ALIGN_PARAGRAPH.CENTER)
    set_docx_text(doc.add_paragraph(), "B2-C1 한글 단어장 ver2", 21, True, "2E74B5", WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(20)
    set_docx_text(p, "영단어 · 품사 · 대표 한글 뜻", 11, False, "4F5D75", WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(140)
    set_docx_text(p, "출처: The Oxford 5000™ (American English)\n원문 목록을 바탕으로 한국어 대표 뜻을 덧붙여 학습용으로 재구성함", 9, False, "687385", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    doc.add_heading("이 단어장 사용법", level=1)
    for text in [
        "각 항목은 영단어, 영어 품사, 대표 한글 뜻 순서입니다.",
        "단어는 알파벳순이며, 매 알파벳을 새 페이지에서 시작합니다.",
        "복수 품사와 동형어는 원문 표기를 보존했습니다. 문맥에 따라 뜻은 달라질 수 있습니다.",
    ]:
        p = doc.add_paragraph(); set_docx_text(p, "• " + text, 10.2)
    groups = defaultdict(list)
    for e in entries: groups[e["word"][0].upper()].append(e)
    letters = sorted(groups)
    doc.add_heading("알파벳 색인", level=1)
    index = doc.add_table(rows=0, cols=4); index.style = "Table Grid"
    for i in range(0, len(letters), 4):
        row = index.add_row()
        for cell, letter in zip(row.cells, letters[i:i + 4]):
            items = groups[letter]; shade(cell, "EAF0F7")
            set_docx_text(cell.paragraphs[0], f"{letter}\n{items[0]['word']} - {items[-1]['word']}\n{len(items)}개", 8.6, True, "163A5F", WD_ALIGN_PARAGRAPH.CENTER)
    geometry(index, [2500] * 4)
    doc.add_page_break()

    for i, letter in enumerate(letters):
        if i: doc.add_page_break()
        doc.add_heading(letter, level=1)
        set_docx_text(doc.add_paragraph(), f"{len(groups[letter])}개 단어", 9.5, False, "687385")
        table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ("영단어", "품사", "대표 한글 뜻")):
            shade(cell, "D9E5F2"); set_docx_text(cell.paragraphs[0], text, 9.2, True, "163A5F", WD_ALIGN_PARAGRAPH.CENTER)
        repeat_header(table.rows[0])
        for e in groups[letter]:
            cells = table.add_row().cells
            set_docx_text(cells[0].paragraphs[0], e["word"], 9.2, True)
            set_docx_text(cells[1].paragraphs[0], e["pos_en"], 8.9, False, "394B63")
            set_docx_text(cells[2].paragraphs[0], e["meaning"], 9.2)
        geometry(table, [3200, 2000, 4800])
    doc.core_properties.title = "Oxford 5000 B2-C1 한글 단어장 ver2"
    doc.core_properties.author = "Codex"
    doc.save(DOCX_PATH)


def make_pdf_styles():
    pdfmetrics.registerFont(TTFont("MalgunGothic", str(FONT)))
    pdfmetrics.registerFont(TTFont("MalgunGothicBold", str(FONT_BOLD)))
    base = getSampleStyleSheet()
    def style(name, font, size, leading, color, align=0, **kw):
        return ParagraphStyle(name, parent=base["Normal"], fontName=font, fontSize=size, leading=leading, textColor=colors.HexColor("#" + color), alignment=align, **kw)
    return {
        "cover_title": style("cover_title", "MalgunGothicBold", 30, 38, "163A5F", TA_CENTER),
        "cover_sub": style("cover_sub", "MalgunGothicBold", 21, 28, "2E74B5", TA_CENTER),
        "note": style("note", "MalgunGothic", 9.5, 15, "687385", TA_CENTER),
        "h1": style("h1", "MalgunGothicBold", 19, 24, "163A5F", 0, spaceAfter=5),
        "body": style("body", "MalgunGothic", 10, 15, "1F2937", 0, spaceAfter=7),
        "small": style("small", "MalgunGothic", 8.5, 11, "687385"),
        "word": style("word", "MalgunGothicBold", 8.9, 11, "1F2937"),
        "cell": style("cell", "MalgunGothic", 8.7, 11, "1F2937"),
        "header": style("header", "MalgunGothicBold", 8.8, 11, "163A5F", TA_CENTER),
        "index": style("index", "MalgunGothicBold", 8.4, 12, "163A5F", TA_CENTER),
    }


def p(text, style):
    return Paragraph(escape(text).replace("\n", "<br/>"), style)


def pdf_footer(canvas, doc):
    canvas.saveState(); canvas.setStrokeColor(colors.HexColor("#D3DAE4")); canvas.setLineWidth(.4); canvas.line(14 * mm, 289 * mm, 196 * mm, 289 * mm)
    canvas.setFillColor(colors.HexColor("#687385")); canvas.setFont("MalgunGothic", 7.5)
    canvas.drawString(14 * mm, 292 * mm, "Oxford 5000 B2-C1 | Korean Vocabulary Book ver2")
    canvas.drawRightString(196 * mm, 8 * mm, f"Page {doc.page}"); canvas.restoreState()


def make_pdf(entries):
    s = make_pdf_styles()
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, leftMargin=14 * mm, rightMargin=14 * mm, topMargin=18 * mm, bottomMargin=16 * mm, title="Oxford 5000 B2-C1 한글 단어장 ver2", author="Codex")
    story = [Spacer(1, 78 * mm), p("Oxford 5000", s["cover_title"]), p("B2-C1 한글 단어장 ver2", s["cover_sub"]), Spacer(1, 9 * mm), p("영단어 · 품사 · 대표 한글 뜻", s["note"]), Spacer(1, 100 * mm), p("출처: The Oxford 5000™ (American English)<br/>원문 목록을 바탕으로 한국어 대표 뜻을 덧붙여 학습용으로 재구성함", s["note"]), PageBreak(), p("이 단어장 사용법", s["h1"])]
    for text in ["각 항목은 영단어, 영어 품사, 대표 한글 뜻 순서입니다.", "단어는 알파벳순이며, 매 알파벳을 새 페이지에서 시작합니다.", "복수 품사와 동형어는 원문 표기를 보존했습니다. 문맥에 따라 뜻은 달라질 수 있습니다."]:
        story.append(p("• " + text, s["body"]))
    groups = defaultdict(list)
    for e in entries: groups[e["word"][0].upper()].append(e)
    letters = sorted(groups)
    story.append(p("알파벳 색인", s["h1"]))
    idx_data = []
    for i in range(0, len(letters), 4):
        row = []
        for letter in letters[i:i + 4]:
            items = groups[letter]; row.append(p(f"{letter}<br/>{items[0]['word']} - {items[-1]['word']}<br/>{len(items)}개", s["index"]))
        row += [""] * (4 - len(row)); idx_data.append(row)
    index = Table(idx_data, colWidths=[45.5 * mm] * 4)
    index.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .35, colors.HexColor("#C5D3E1")), ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EAF0F7")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
    story += [index, PageBreak()]
    for i, letter in enumerate(letters):
        if i: story.append(PageBreak())
        story += [p(letter, s["h1"]), p(f"{len(groups[letter])}개 단어", s["small"])]
        data = [[p("영단어", s["header"]), p("품사", s["header"]), p("대표 한글 뜻", s["header"])]]
        for e in groups[letter]: data.append([p(e["word"], s["word"]), p(e["pos_en"], s["cell"]), p(e["meaning"], s["cell"])])
        table = Table(data, colWidths=[60 * mm, 38 * mm, 84 * mm], repeatRows=1, splitByRow=1, hAlign="LEFT")
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), .25, colors.HexColor("#C7D2DE")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E5F2")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 3.5), ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5), ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4)]))
        story.append(table)
    doc.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)


if __name__ == "__main__":
    OUT.mkdir(exist_ok=True)
    data = load_entries()
    make_docx(data)
    make_pdf(data)
    print(f"Created {DOCX_PATH}\nCreated {PDF_PATH}\nEntries: {len(data)}")
