from __future__ import annotations

import csv
import json
import re
import sys
import time
import urllib.parse
import urllib.request
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from pypdf import PdfReader


SOURCE_PDF = Path(r"C:\Users\dhth0\Downloads\American_Oxford_5000.pdf")
ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "output"
DATA = ROOT / "build"
DOCX_PATH = OUTPUT / "Oxford_5000_B2-C1_Korean_Vocabulary.docx"
CSV_PATH = DATA / "oxford_5000_entries.csv"
REPORT_PATH = DATA / "validation_report.json"

POS_RE = r"(?:n\.|v\.|adj\.|adv\.|prep\.|conj\.|pron\.|det\.|exclam\.|number)"
ENTRY_RE = re.compile(
    rf"^(?P<word>.+?)\s+(?P<detail>(?:(?:{POS_RE})(?:,\s*(?:{POS_RE}))*\s*(?:B1|B2|C1)(?:,\s*)?)+)$"
)
DETAIL_RE = re.compile(
    rf"(?P<poses>(?:{POS_RE})(?:,\s*(?:{POS_RE}))*)\s*(?P<level>B1|B2|C1)"
)
POS_KO = {
    "n.": "명사", "v.": "동사", "adj.": "형용사", "adv.": "부사",
    "prep.": "전치사", "conj.": "접속사", "pron.": "대명사", "det.": "한정사", "exclam.": "감탄사", "number": "수사",
}

# These source labels deliberately distinguish senses that automated translation tends to blur.
GLOSS_OVERRIDES = {
    "grave (for dead person)": "무덤",
    "grave (serious)": "심각한, 중대한",
    "strip (long narrow piece)": "가늘고 긴 조각, 띠",
    "strip (remove clothes/a layer)": "벗기다; 벗다",
    "minute2": "아주 작은, 미세한",
    "recount1": "이야기를 들려주다, 자세히 설명하다",
    "scratch": "긁다; 긁힌 자국",
    "seal": "봉인하다; 봉인, 도장",
    "dynamic": "역동적인; 동력",
    "acid": "산; 신맛의",
    "alert": "경계하다; 경보; 기민한",
    "alien": "외국인; 낯선",
    "abuse": "학대, 남용; 학대하다, 남용하다",
    "advocate": "옹호자; 옹호하다",
    "combat": "전투; 싸우다",
    "ease": "수월함; 완화하다",
    "grasp": "움켜쥠; 이해하다, 움켜쥐다",
    "grin": "활짝 웃다; 활짝 웃음",
    "grip": "움켜쥠; 꽉 잡다",
    "harvest": "수확; 수확하다",
    "mature": "성숙한; 성숙하다",
    "principal": "주요한; 교장, 원금",
    "probe": "탐사, 조사; 조사하다",
    "raid": "급습; 급습하다",
    "rally": "집회; 다시 회복하다, 결집하다",
    "spin": "회전; 회전하다",
    "span": "범위; 걸치다",
    "spare": "여분의; 아끼다, 용서하다",
    "warrant": "영장, 보증; 보증하다",
    "worship": "숭배; 숭배하다",
    "yield": "산출량; 양보하다, 생산하다",
}


def normalize_source_lines() -> list[str]:
    reader = PdfReader(str(SOURCE_PDF))
    lines: list[str] = []
    for page in reader.pages:
        lines.extend((page.extract_text() or "").splitlines())

    clean = []
    for raw in lines:
        line = raw.strip()
        if not line or line.startswith("© Oxford") or line.startswith("The Oxford 5000"):
            continue
        if line.startswith("3000, it includes an additional 2000 words"):
            continue
        clean.append(line)

    merged: list[str] = []
    i = 0
    while i < len(clean):
        line = clean[i]
        nxt = clean[i + 1] if i + 1 < len(clean) else ""
        # The source uses superscripts and line wrapping in these six entries.
        if line == "recount" and re.match(r"^1\s+v\.\s+C1$", nxt):
            merged.append("recount1 v. C1")
            i += 2
            continue
        if line in {"sacred", "seal", "strip (remove clothes/a layer)"} and re.match(rf"^{POS_RE}", nxt):
            merged.append(f"{line} {nxt}")
            i += 2
            continue
        if line == "scratch v., n. B" and nxt == "2":
            merged.append("scratch v., n. B2")
            i += 2
            continue
        merged.append(line)
        i += 1

    # A few entries omit the visual space before the CEFR level in extraction.
    normalized = []
    for line in merged:
        line = line.replace("adj./adv.", "adj., adv.")
        line = re.sub(r"(?<=[.])(?=(?:B1|B2|C1)$)", " ", line)
        line = re.sub(r"(?<=\.),\s+(?=(?:B1|B2|C1)$)", " ", line)
        normalized.append(line)
    return normalized


def detail_to_korean(detail: str) -> tuple[str, str]:
    groups = []
    levels = []
    for match in DETAIL_RE.finditer(detail):
        poses = [POS_KO[p] for p in re.findall(POS_RE, match.group("poses"))]
        groups.append("·".join(poses))
        levels.append(match.group("level"))
    if not groups:
        raise ValueError(f"Could not parse part of speech: {detail}")
    return " / ".join(groups), " / ".join(levels)


def parse_entries() -> list[dict[str, str]]:
    entries = []
    bad = []
    for line in normalize_source_lines():
        match = ENTRY_RE.match(line)
        if not match:
            bad.append(line)
            continue
        word = match.group("word")
        detail = match.group("detail")
        pos_ko, level = detail_to_korean(detail)
        entries.append({"word": word, "source_detail": detail, "pos": pos_ko, "level": level})
    if bad:
        raise ValueError(f"Unparsed source lines ({len(bad)}): {bad}")
    return entries


def translate_batch(terms: list[str]) -> list[str]:
    text = "\n".join(terms)
    query = urllib.parse.urlencode({"client": "gtx", "sl": "en", "tl": "ko", "dt": "t", "q": text})
    request = urllib.request.Request(
        "https://translate.googleapis.com/translate_a/single?" + query,
        headers={"User-Agent": "Mozilla/5.0"},
    )
    last_error = None
    for attempt in range(4):
        try:
            with urllib.request.urlopen(request, timeout=45) as response:
                payload = json.loads(response.read().decode("utf-8"))
            translated = "".join(part[0] for part in payload[0]).strip()
            result = translated.split("\n")
            if len(result) != len(terms):
                raise ValueError(f"Expected {len(terms)} translations, received {len(result)}")
            return [value.strip() for value in result]
        except Exception as exc:  # Transient gateway/rate-limit errors are retried.
            last_error = exc
            time.sleep(2 * (attempt + 1))
    raise RuntimeError(f"Translation request failed: {last_error}")


def add_translations(entries: list[dict[str, str]]) -> None:
    # Keep each URL comfortably below common proxy and server limits.
    batches: list[list[dict[str, str]]] = []
    current: list[dict[str, str]] = []
    current_len = 0
    for entry in entries:
        size = len(entry["word"]) + 1
        if current and current_len + size > 2400:
            batches.append(current)
            current, current_len = [], 0
        current.append(entry)
        current_len += size
    if current:
        batches.append(current)

    for number, batch in enumerate(batches, start=1):
        translations = translate_batch([entry["word"] for entry in batch])
        for entry, korean in zip(batch, translations):
            entry["meaning"] = GLOSS_OVERRIDES.get(entry["word"], korean)
        print(f"Translated batch {number}/{len(batches)} ({len(batch)} entries)", flush=True)
        time.sleep(0.25)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Malgun Gothic"
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_text(paragraph, text: str, size: float = 9.2, bold: bool = False, color: str | None = None, align=None) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    if align is not None:
        paragraph.alignment = align


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def make_document(entries: list[dict[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color in (("Title", 26, "163A5F"), ("Heading 1", 19, "163A5F"), ("Heading 2", 13, "2E74B5")):
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header_p = section.header.paragraphs[0]
    set_paragraph_text(header_p, "Oxford 5000 B2-C1 | Korean Vocabulary Book", 8.5, color="687385")
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_text(footer_p, "", 8.5, color="687385", align=WD_ALIGN_PARAGRAPH.RIGHT)
    footer_p.add_run("Page ")
    add_page_field(footer_p)

    # Cover
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(124)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(p, "Oxford 5000", 30, bold=True, color="163A5F", align=WD_ALIGN_PARAGRAPH.CENTER)
    p = document.add_paragraph()
    set_paragraph_text(p, "B2-C1 한글 단어장", 21, bold=True, color="2E74B5", align=WD_ALIGN_PARAGRAPH.CENTER)
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    set_paragraph_text(p, "American English | 약 2,000개 고급 학습 단어", 11, color="4F5D75", align=WD_ALIGN_PARAGRAPH.CENTER)
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(140)
    set_paragraph_text(p, "출처: The Oxford 5000™ (American English)\n원문 목록을 바탕으로 한국어 대표 뜻을 덧붙여 학습용으로 재구성함", 9, color="687385", align=WD_ALIGN_PARAGRAPH.CENTER)
    document.add_page_break()

    # Guide
    document.add_heading("이 단어장 사용법", level=1)
    guide = [
        "각 항목은 영단어, 원문의 품사, CEFR 레벨, 대표 한글 뜻 순서입니다.",
        "단어는 알파벳순이며, 매 알파벳을 새 페이지에서 시작해 빠르게 찾아볼 수 있습니다.",
        "복수 품사와 동형어는 원문 표기를 보존했습니다. 문맥에 따라 뜻은 달라질 수 있습니다.",
        "레벨은 원문 표기를 그대로 따릅니다. 원문에 포함된 B1 표기도 수정하지 않았습니다.",
    ]
    for item in guide:
        p = document.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Mm(4)
        p.paragraph_format.first_line_indent = Mm(-4)
        set_paragraph_text(p, "• " + item, 10.2)

    document.add_paragraph()
    document.add_heading("알파벳 색인", level=1)
    groups: dict[str, list[dict[str, str]]] = defaultdict(list)
    for entry in entries:
        groups[entry["word"][0].upper()].append(entry)
    index_table = document.add_table(rows=0, cols=4)
    index_table.style = "Table Grid"
    index_rows = []
    letters = sorted(groups)
    for start in range(0, len(letters), 4):
        index_rows.append(letters[start:start + 4])
    for row_letters in index_rows:
        row = index_table.add_row()
        for cell, letter in zip(row.cells, row_letters):
            values = groups[letter]
            set_cell_shading(cell, "EAF0F7")
            set_paragraph_text(cell.paragraphs[0], f"{letter}\n{values[0]['word']} - {values[-1]['word']}\n{len(values)}개", 8.6, bold=True, color="163A5F", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(index_table, [2500, 2500, 2500, 2500])
    document.add_page_break()

    table_widths = [3000, 2100, 1100, 3800]
    for position, letter in enumerate(letters):
        if position:
            document.add_page_break()
        document.add_heading(letter, level=1)
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        set_paragraph_text(p, f"{len(groups[letter])}개 단어", 9.5, color="687385")
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        header = table.rows[0]
        for cell, text in zip(header.cells, ("영단어", "품사", "레벨", "대표 한글 뜻")):
            set_cell_shading(cell, "D9E5F2")
            set_paragraph_text(cell.paragraphs[0], text, 9.2, bold=True, color="163A5F", align=WD_ALIGN_PARAGRAPH.CENTER)
        repeat_header(header)
        for entry in groups[letter]:
            cells = table.add_row().cells
            set_paragraph_text(cells[0].paragraphs[0], entry["word"], 9.2, bold=True)
            set_paragraph_text(cells[1].paragraphs[0], entry["pos"], 8.7, color="394B63")
            set_paragraph_text(cells[2].paragraphs[0], entry["level"], 8.7, bold=True, color="2E74B5", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_paragraph_text(cells[3].paragraphs[0], entry["meaning"], 9.2)
        set_table_geometry(table, table_widths)

    document.core_properties.title = "Oxford 5000 B2-C1 한글 단어장"
    document.core_properties.subject = "American English Oxford 5000 vocabulary with Korean glosses"
    document.core_properties.author = "Codex"
    document.save(DOCX_PATH)


def write_data(entries: list[dict[str, str]]) -> None:
    with CSV_PATH.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=["word", "source_detail", "pos", "level", "meaning"])
        writer.writeheader()
        writer.writerows(entries)
    report = {
        "source_pdf": str(SOURCE_PDF),
        "entry_count": len(entries),
        "missing_meanings": sum(not x.get("meaning") for x in entries),
        "letters": {letter: sum(x["word"].startswith(letter) for x in entries) for letter in "ABCDEFGHIJKLMNOPQRSTUVWXYZ" if any(x["word"].startswith(letter) for x in entries)},
    }
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


def main() -> int:
    if not SOURCE_PDF.exists():
        raise FileNotFoundError(SOURCE_PDF)
    OUTPUT.mkdir(exist_ok=True)
    DATA.mkdir(exist_ok=True)
    entries = parse_entries()
    print(f"Parsed {len(entries)} source entries", flush=True)
    add_translations(entries)
    if any(not entry.get("meaning") for entry in entries):
        raise ValueError("A Korean gloss is missing")
    write_data(entries)
    make_document(entries)
    print(DOCX_PATH)
    return 0


if __name__ == "__main__":
    sys.exit(main())
